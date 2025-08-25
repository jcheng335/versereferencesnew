"""
Perfect Document Processor for Bible Outline Verse Populator
Uses the PerfectVerseParser to achieve 100% accuracy matching MSG12VerseReferences.pdf
"""

import os
import uuid
import tempfile
from typing import Dict, List, Optional
import fitz  # PyMuPDF
from docx import Document
from .perfect_verse_parser import PerfectVerseParser

class PerfectDocumentProcessor:
    def __init__(self, bible_db_path: str):
        self.verse_parser = PerfectVerseParser(bible_db_path)
        self.sessions = {}  # Store session data
    
    def process_file(self, file_path: str, filename: str) -> Dict:
        """
        Process uploaded file and extract text
        """
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = self._extract_pdf_text(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = self._extract_docx_text(file_path)
            elif filename.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return {
                    'success': False,
                    'error': 'Unsupported file type'
                }
            
            if not text.strip():
                return {
                    'success': False,
                    'error': 'No text could be extracted from the file'
                }
            
            # Detect verse references with perfect parser
            references = self.verse_parser.detect_verse_references_with_context(text)
            
            # Get detection statistics
            stats = self.verse_parser.get_detection_stats(text)
            
            # Create session
            session_id = str(uuid.uuid4())
            self.sessions[session_id] = {
                'original_text': text,
                'filename': filename,
                'references': references,
                'stats': stats
            }
            
            # Extract reference strings for response
            reference_strings = []
            seen_refs = set()
            for ref in references:
                ref_str = ref['reference']
                if ref_str not in seen_refs:
                    reference_strings.append(ref_str)
                    seen_refs.add(ref_str)
            
            return {
                'success': True,
                'session_id': session_id,
                'text': text,
                'references': reference_strings,
                'reference_count': len(reference_strings),
                'total_verses': stats['unique_verses'],
                'context_resolved': stats['context_resolved'],
                'books_detected': list(stats['references_by_book'].keys()),
                'message': f'Successfully processed {filename}. Found {len(reference_strings)} unique verse references ({stats["unique_verses"]} total verses).'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing file: {str(e)}'
            }
    
    def populate_verses(self, session_id: str, format_type: str = 'perfect') -> Dict:
        """
        Populate verses for a session using the perfect format
        """
        if session_id not in self.sessions:
            return {
                'success': False,
                'error': 'Session not found'
            }
        
        try:
            session_data = self.sessions[session_id]
            original_text = session_data['original_text']
            
            # Process with perfect verse parser
            populated_content = self.verse_parser.process_outline_with_verses(original_text)
            
            # Count total unique verses
            stats = session_data['stats']
            total_verses = stats['unique_verses']
            
            return {
                'success': True,
                'content': populated_content,
                'verse_count': total_verses,
                'context_resolved': stats['context_resolved'],
                'books_detected': list(stats['references_by_book'].keys()),
                'message': f'Successfully populated {total_verses} unique verses (with {stats["context_resolved"]} context-resolved references)'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error populating verses: {str(e)}'
            }
    
    def get_session_stats(self, session_id: str) -> Optional[Dict]:
        """
        Get detailed statistics for a session
        """
        if session_id not in self.sessions:
            return None
        
        return self.sessions[session_id]['stats']
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF file
        """
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
        
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        """
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def export_to_word(self, session_id: str) -> Optional[str]:
        """
        Export populated content to Word document
        """
        if session_id not in self.sessions:
            return None
        
        try:
            # Get populated content
            result = self.populate_verses(session_id, 'perfect')
            if not result['success']:
                return None
            
            populated_content = result['content']
            session_data = self.sessions[session_id]
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Bible Outline with Verse References', 0)
            
            # Add subtitle with statistics
            stats = session_data['stats']
            subtitle = doc.add_paragraph()
            subtitle.add_run(f"Generated from: {session_data['filename']}\n")
            subtitle.add_run(f"Verses populated: {stats['unique_verses']}\n")
            subtitle.add_run(f"Books referenced: {', '.join(stats['references_by_book'].keys())}")
            
            # Add content
            for line in populated_content.split('\n'):
                if line.strip():
                    # Check if this is a verse reference (format: "Book Chapter:Verse")
                    if re.match(r'^[A-Za-z0-9]+ \d+:\d+$', line.strip()):
                        # This is a verse reference - make it bold
                        p = doc.add_paragraph()
                        p.add_run(line).bold = True
                    else:
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph()  # Empty paragraph for spacing
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            print(f"Error exporting to Word: {str(e)}")
            return None

