"""
Accurate Inline Processor - Fixes verse placement and formatting issues
Only inserts verses where specifically referenced, preserves original structure
"""

import os
import uuid
import pdfplumber
from docx import Document
import tempfile
import re
from typing import Dict, List, Optional, Any
from .sqlite_bible_database import SQLiteBibleDatabase

class AccurateInlineProcessor:
    def __init__(self, db_path: str):
        self.db = SQLiteBibleDatabase(db_path)
        self.sessions = {}
        
        # Book abbreviation mapping
        self.book_abbreviations = {
            'Gen': 'Genesis', 'Exo': 'Exodus', 'Lev': 'Leviticus', 'Num': 'Numbers', 'Deut': 'Deuteronomy',
            'Josh': 'Joshua', 'Judg': 'Judges', 'Ruth': 'Ruth', '1 Sam': '1 Samuel', '2 Sam': '2 Samuel',
            '1 Kings': '1 Kings', '2 Kings': '2 Kings', '1 Chron': '1 Chronicles', '2 Chron': '2 Chronicles',
            'Ezra': 'Ezra', 'Neh': 'Nehemiah', 'Esth': 'Esther', 'Job': 'Job', 'Psa': 'Psalms',
            'Prov': 'Proverbs', 'Eccl': 'Ecclesiastes', 'S.S.': 'Song of Songs', 'Isa': 'Isaiah',
            'Jer': 'Jeremiah', 'Lam': 'Lamentations', 'Ezek': 'Ezekiel', 'Dan': 'Daniel',
            'Hosea': 'Hosea', 'Joel': 'Joel', 'Amos': 'Amos', 'Obad': 'Obadiah', 'Jonah': 'Jonah',
            'Micah': 'Micah', 'Nahum': 'Nahum', 'Hab': 'Habakkuk', 'Zeph': 'Zephaniah', 'Hag': 'Haggai',
            'Zech': 'Zechariah', 'Mal': 'Malachi',
            'Matt': 'Matthew', 'Mark': 'Mark', 'Luke': 'Luke', 'John': 'John', 'Acts': 'Acts',
            'Rom': 'Romans', '1 Cor': '1 Corinthians', '2 Cor': '2 Corinthians', 'Gal': 'Galatians',
            'Eph': 'Ephesians', 'Phil': 'Philippians', 'Col': 'Colossians', '1 Thes': '1 Thessalonians',
            '2 Thes': '2 Thessalonians', '1 Tim': '1 Timothy', '2 Tim': '2 Timothy', 'Titus': 'Titus',
            'Philem': 'Philemon', 'Heb': 'Hebrews', 'James': 'James', '1 Pet': '1 Peter', '2 Pet': '2 Peter',
            '1 John': '1 John', '2 John': '2 John', '3 John': '3 John', 'Jude': 'Jude', 'Rev': 'Revelation'
        }
    
    def process_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Process uploaded file with accurate verse detection
        """
        try:
            # Generate session ID
            session_id = str(uuid.uuid4())
            
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                content = self._extract_pdf_text(file_path)
            elif filename.lower().endswith(('.docx', '.doc')):
                content = self._extract_docx_text(file_path)
            else:
                # Plain text
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # Find only inline verse references (not Scripture Reading sections)
            references = self._find_inline_references(content)
            
            # Store session data
            self.sessions[session_id] = {
                'original_content': content,
                'original_filename': filename,
                'references': references,
                'populated_content': None
            }
            
            return {
                'success': True,
                'session_id': session_id,
                'content': content,
                'references_found': len(references),
                'total_verses': len(references),  # Each reference = 1 verse for inline
                'filename': filename
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
        
        return self._clean_extracted_text(text)
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return self._clean_extracted_text(text)
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace but preserve structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive spaces within lines
            cleaned_line = re.sub(r' +', ' ', line.strip())
            cleaned_lines.append(cleaned_line)
        
        # Remove excessive empty lines but preserve paragraph breaks
        result_lines = []
        empty_count = 0
        
        for line in cleaned_lines:
            if not line:
                empty_count += 1
                if empty_count <= 2:  # Allow max 2 consecutive empty lines
                    result_lines.append(line)
            else:
                empty_count = 0
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def _find_inline_references(self, text: str) -> List[Dict]:
        """
        Find only inline verse references, NOT Scripture Reading sections
        """
        references = []
        lines = text.split('\n')
        current_book = None
        current_chapter = None
        
        for line_idx, line in enumerate(lines):
            # Skip Scripture Reading lines - these are for reference only
            if re.search(r'Scripture Reading:', line, re.IGNORECASE):
                continue
            
            # Find inline verse references in outline points
            inline_refs = self._find_inline_refs_in_line(line, line_idx, current_book, current_chapter)
            
            # Update context from full references
            for ref in inline_refs:
                if ref.get('book') and ref.get('chapter'):
                    current_book = ref['book']
                    current_chapter = ref['chapter']
            
            references.extend(inline_refs)
        
        return references
    
    def _find_inline_refs_in_line(self, line: str, line_idx: int, current_book: str, current_chapter: int) -> List[Dict]:
        """
        Find inline verse references in a single line
        """
        references = []
        
        # Pattern 1: Full references like "1 Cor. 12:14", "Eph. 4:7"
        full_ref_pattern = r'(?:cf\.\s+)?([123]?\s*[A-Za-z]+\.?)\s+(\d+):(\d+)(?:-(\d+))?'
        for match in re.finditer(full_ref_pattern, line):
            book_abbrev = match.group(1).strip().replace('.', '')
            book_name = self.book_abbreviations.get(book_abbrev, book_abbrev)
            chapter = int(match.group(2))
            start_verse = int(match.group(3))
            end_verse = int(match.group(4)) if match.group(4) else start_verse
            
            # Only add if this is in an outline point, not a Scripture Reading
            if self._is_outline_point(line):
                references.append({
                    'type': 'full_reference',
                    'book': book_name,
                    'chapter': chapter,
                    'start_verse': start_verse,
                    'end_verse': end_verse,
                    'line_idx': line_idx,
                    'position': match.span(),
                    'original_text': match.group(0)
                })
        
        # Pattern 2: Verse-only references like "v. 7", "vv. 15-16"
        verse_only_pattern = r'v{1,2}\.\s*(\d+)(?:-(\d+))?'
        for match in re.finditer(verse_only_pattern, line):
            start_verse = int(match.group(1))
            end_verse = int(match.group(2)) if match.group(2) else start_verse
            
            # Only add if we have context and this is in an outline point
            if current_book and current_chapter and self._is_outline_point(line):
                references.append({
                    'type': 'verse_only',
                    'book': current_book,
                    'chapter': current_chapter,
                    'start_verse': start_verse,
                    'end_verse': end_verse,
                    'line_idx': line_idx,
                    'position': match.span(),
                    'original_text': match.group(0)
                })
        
        return references
    
    def _is_outline_point(self, line: str) -> bool:
        """
        Check if line is an outline point (not a Scripture Reading section)
        """
        line = line.strip()
        
        # Skip Scripture Reading lines
        if re.search(r'Scripture Reading:', line, re.IGNORECASE):
            return False
        
        # Check for outline markers
        outline_patterns = [
            r'^[IVX]+\.',  # Roman numerals
            r'^[A-Z]\.',   # Capital letters
            r'^[0-9]+\.',  # Numbers
            r'^[a-z]\.',   # Lowercase letters
        ]
        
        for pattern in outline_patterns:
            if re.match(pattern, line):
                return True
        
        # Also include lines that contain verse references but aren't Scripture Reading
        if re.search(r'v{1,2}\.\s*\d+|[A-Za-z]+\.\s*\d+:\d+', line):
            return True
        
        return False
    
    def populate_verses(self, session_id: str, format_type: str = 'inline') -> Dict[str, Any]:
        """
        Populate verses with accurate inline placement
        """
        if session_id not in self.sessions:
            return {
                'success': False,
                'error': 'Session not found'
            }
        
        try:
            session_data = self.sessions[session_id]
            original_content = session_data['original_content']
            references = session_data['references']
            
            # Process content line by line with accurate verse insertion
            populated_content = self._insert_verses_accurately(original_content, references)
            
            # Update session
            session_data['populated_content'] = populated_content
            
            return {
                'success': True,
                'populated_content': populated_content,
                'verse_count': len(references),
                'context_resolved': sum(1 for ref in references if ref.get('type') == 'verse_only'),
                'format': format_type,
                'message': f'Successfully populated {len(references)} verses inline with accurate placement'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error populating verses: {str(e)}'
            }
    
    def _insert_verses_accurately(self, content: str, references: List[Dict]) -> str:
        """
        Insert verses with accurate placement and perfect formatting preservation
        """
        lines = content.split('\n')
        result_lines = []
        
        # Group references by line
        refs_by_line = {}
        for ref in references:
            line_idx = ref['line_idx']
            if line_idx not in refs_by_line:
                refs_by_line[line_idx] = []
            refs_by_line[line_idx].append(ref)
        
        for line_idx, line in enumerate(lines):
            # Preserve exact original line with proper formatting
            formatted_line = self._preserve_outline_formatting(line)
            result_lines.append(formatted_line)
            
            # Add verses for this line if any
            if line_idx in refs_by_line:
                # Add verses immediately after the reference line
                for ref in refs_by_line[line_idx]:
                    # Get verse text for each verse in the range
                    for verse_num in range(ref['start_verse'], ref['end_verse'] + 1):
                        verse_data = self.db.lookup_verse(ref['book'], ref['chapter'], verse_num)
                        if verse_data and 'text' in verse_data:
                            # Format verse reference and text with proper indentation
                            book_abbrev = self._get_book_abbreviation(ref['book'])
                            
                            # Create verse reference line
                            ref_line = f"<span class='verse-ref'>{book_abbrev} {ref['chapter']}:{verse_num}</span>"
                            result_lines.append(ref_line)
                            
                            # Create verse text line with proper indentation
                            verse_text = f"<span class='verse-text'>{verse_data['text']}</span>"
                            result_lines.append(verse_text)
                
                # Add empty line after verses for readability
                result_lines.append("")
        
        return '\n'.join(result_lines)
    
    def _preserve_outline_formatting(self, line: str) -> str:
        """
        Preserve exact outline formatting with proper HTML structure
        """
        if not line.strip():
            return line
        
        # Detect indentation level
        original_indent = len(line) - len(line.lstrip())
        indent_chars = line[:original_indent]
        line_content = line[original_indent:]
        
        # Apply formatting based on outline structure
        formatted_content = self._apply_outline_styling(line_content)
        
        # Reconstruct with original indentation preserved
        return indent_chars + formatted_content
    
    def _apply_outline_styling(self, line_content: str) -> str:
        """
        Apply proper styling to outline elements
        """
        stripped = line_content.strip()
        
        # Roman numerals (I., II., III., IV., etc.) - Bold and prominent
        if re.match(r'^[IVX]+\.', stripped):
            return f"<strong class='roman-numeral'>{line_content}</strong>"
        
        # Capital letters (A., B., C., etc.) - Bold
        elif re.match(r'^[A-Z]\. ', stripped):
            return f"<strong class='letter-point'>{line_content}</strong>"
        
        # Numbers (1., 2., 3., etc.) - Bold
        elif re.match(r'^\d+\. ', stripped):
            return f"<strong class='number-point'>{line_content}</strong>"
        
        # Lowercase letters (a., b., c., etc.) - Bold
        elif re.match(r'^[a-z]\. ', stripped):
            return f"<strong class='sub-letter-point'>{line_content}</strong>"
        
        # Scripture Reading section
        elif 'Scripture Reading:' in line_content:
            return f"<div class='scripture-reading'><strong>{line_content}</strong></div>"
        
        # Message titles
        elif line_content.startswith('Message '):
            return f"<h2 class='message-title'>{line_content}</h2>"
        
        # General subject
        elif line_content.isupper() and len(line_content) > 10:
            return f"<h3 class='general-subject'>{line_content}</h3>"
        
        # Regular content
        else:
            return line_content
    def _format_outline_line(self, line: str) -> str:
        """
        Format outline lines with proper HTML formatting
        """
        if not line.strip():
            return line
        
        # Format Roman numerals
        if re.match(r'^\s*[IVX]+\.', line.strip()):
            return f"<strong class='roman-numeral'>{line}</strong>"
        
        # Format capital letter points
        if re.match(r'^\s*[A-Z]\.', line.strip()):
            return f"<strong class='outline-point'>{line}</strong>"
        
        # Format numbered points
        if re.match(r'^\s*\d+\.', line.strip()):
            return f"<strong class='numbered-point'>{line}</strong>"
        
        return line
    
    def _get_book_abbreviation(self, book_name: str) -> str:
        """Get standard abbreviation for book name"""
        abbrev_map = {v: k for k, v in self.book_abbreviations.items()}
        return abbrev_map.get(book_name, book_name)
    
    def export_clean_text(self, session_id: str) -> Optional[str]:
        """Export clean text without HTML tags for OneNote"""
        if session_id not in self.sessions:
            return None
        
        populated_content = self.sessions[session_id].get('populated_content')
        if not populated_content:
            return None
        
        # Remove HTML tags but keep content
        clean_text = populated_content
        
        # Remove verse span tags but keep content
        clean_text = re.sub(r'<span class=[\'"]verse-ref[\'"]>(.*?)</span>', r'\1', clean_text)
        clean_text = re.sub(r'<span class=[\'"]verse-text[\'"]>(.*?)</span>', r'\1', clean_text)
        
        # Remove other HTML tags but keep content
        clean_text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'\1', clean_text)
        clean_text = re.sub(r'<em>(.*?)</em>', r'\1', clean_text)
        clean_text = re.sub(r'<[^>]*>', '', clean_text)
        
        # Clean up whitespace
        clean_text = re.sub(r'\n\s*\n\s*\n', '\n\n', clean_text)
        
        return clean_text.strip()
    
    def export_to_word(self, session_id: str) -> Optional[str]:
        """Export to Word document"""
        if session_id not in self.sessions:
            return None
        
        session_data = self.sessions[session_id]
        populated_content = session_data.get('populated_content')
        
        if not populated_content:
            return None
        
        try:
            doc = Document()
            
            # Process content with proper formatting
            for line in populated_content.split('\n'):
                if line.strip():
                    # Check if this is a verse reference or verse text (HTML spans)
                    if '<span class="verse-ref">' in line:
                        # Extract verse reference and make it bold and blue
                        ref_text = line.replace('<span class="verse-ref">', '').replace('</span>', '')
                        p = doc.add_paragraph()
                        run = p.add_run(ref_text)
                        run.bold = True
                    elif '<span class="verse-text">' in line:
                        # Extract verse text
                        verse_text = line.replace('<span class="verse-text">', '').replace('</span>', '')
                        p = doc.add_paragraph()
                        p.add_run(verse_text)
                    else:
                        # Regular outline text
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph()  # Empty paragraph for spacing
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            print(f"Error exporting to Word: {str(e)}")
            return None

