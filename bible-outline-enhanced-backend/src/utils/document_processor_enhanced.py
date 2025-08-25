import re
import os
import uuid
import tempfile
from typing import List, Dict, Tuple, Optional
from docx import Document
from PyPDF2 import PdfReader
from .enhanced_verse_parser import EnhancedVerseParser
from .sqlite_bible_database import SQLiteBibleDatabase

class EnhancedDocumentProcessor:
    def __init__(self):
        self.db_path = os.path.join(os.path.dirname(__file__), '..', '..', 'bible_verses.db')
        self.bible_db = SQLiteBibleDatabase(self.db_path)
        self.verse_parser = EnhancedVerseParser()
        self.sessions = {}  # In-memory session storage
    
    def process_uploaded_file(self, file_path: str, filename: str, text_content: str = None) -> Dict:
        """Process uploaded file and extract verse references"""
        try:
            # Extract text from file or use provided text
            if text_content:
                text = text_content
            else:
                text = self._extract_text_from_file(file_path, filename)
            
            if not text:
                return {'success': False, 'error': 'Could not extract text from file'}
            
            # Extract verse references using enhanced parser
            references = self.verse_parser.extract_references_from_text(text)
            
            # Create session
            session_id = str(uuid.uuid4())
            self.sessions[session_id] = {
                'original_text': text,
                'filename': filename,
                'references': references,
                'processed_content': text,
                'created_at': str(uuid.uuid4())  # Placeholder timestamp
            }
            
            return {
                'success': True,
                'session_id': session_id,
                'text': text,
                'references': references,
                'reference_count': len(references),
                'message': f'Successfully processed {filename}. Found {len(references)} verse references.'
            }
        
        except Exception as e:
            return {'success': False, 'error': f'Error processing file: {str(e)}'}
    
    def _extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Extract text from PDF, Word document, or text file"""
        try:
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['doc', 'docx']:
                return self._extract_from_docx(file_path)
            elif file_ext == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        
        except Exception as e:
            raise Exception(f"Error extracting text from {filename}: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")
    
    def get_session(self, session_id: str) -> Optional[Dict]:
        """Get session data"""
        return self.sessions.get(session_id)
    
    def update_session_content(self, session_id: str, content: str) -> bool:
        """Update session content"""
        if session_id in self.sessions:
            self.sessions[session_id]['processed_content'] = content
            return True
        return False
    
    def populate_document_with_verses(self, session_id: str, format_type: str = "bottom") -> Dict:
        """Populate document with Bible verses"""
        try:
            session = self.get_session(session_id)
            if not session:
                return {'success': False, 'error': 'Session not found'}
            
            references = session['references']
            populated_verses = []
            
            # Process each reference
            for ref_string in references:
                parsed_refs = self.verse_parser.parse_reference(ref_string)
                
                for parsed_ref in parsed_refs:
                    # Look up verse in database
                    verse_data = self.bible_db.lookup_verse(
                        parsed_ref['book'], 
                        parsed_ref['chapter'], 
                        parsed_ref['verse']
                    )
                    
                    if verse_data:
                        populated_verses.append({
                            'reference': parsed_ref['reference'],
                            'text': verse_data['text'],
                            'book': parsed_ref['book'],
                            'chapter': parsed_ref['chapter'],
                            'verse': parsed_ref['verse']
                        })
            
            # Format output based on user preference
            if format_type == "bottom":
                formatted_content = self._format_verses_at_bottom(
                    session['processed_content'], 
                    populated_verses
                )
            else:
                formatted_content = self._format_verses_inline(
                    session['processed_content'], 
                    populated_verses
                )
            
            # Update session
            self.sessions[session_id]['populated_content'] = formatted_content
            self.sessions[session_id]['populated_verses'] = populated_verses
            
            return {
                'success': True,
                'content': formatted_content,
                'verses': populated_verses,
                'verse_count': len(populated_verses),
                'message': f'Successfully populated {len(populated_verses)} verses'
            }
        
        except Exception as e:
            return {'success': False, 'error': f'Error populating verses: {str(e)}'}
    
    def _format_verses_at_bottom(self, content: str, verses: List[Dict]) -> str:
        """Format verses at the bottom of the document"""
        if not verses:
            return content
        
        # Add separator and verses at the bottom
        formatted_content = content + "\\n\\n" + "="*50 + "\\n"
        formatted_content += "BIBLE VERSES\\n"
        formatted_content += "="*50 + "\\n\\n"
        
        # Group verses by reference for clean display
        verse_dict = {}
        for verse in verses:
            ref = verse['reference']
            if ref not in verse_dict:
                verse_dict[ref] = verse['text']
        
        # Add each unique verse
        for reference, text in verse_dict.items():
            formatted_content += f"{reference}\\n"
            formatted_content += f"{text}\\n\\n"
        
        return formatted_content
    
    def _format_verses_inline(self, content: str, verses: List[Dict]) -> str:
        """Format verses inline with the content"""
        if not verses:
            return content
        
        formatted_content = content
        
        # Create a mapping of references to verse text
        verse_dict = {}
        for verse in verses:
            ref = verse['reference']
            if ref not in verse_dict:
                verse_dict[ref] = verse['text']
        
        # Replace references with reference + verse text
        for reference, text in verse_dict.items():
            # Find the reference in the content and add the verse text after it
            pattern = re.escape(reference)
            replacement = f"{reference}\\n{text}\\n"
            formatted_content = re.sub(pattern, replacement, formatted_content)
        
        return formatted_content
    
    def export_to_docx(self, session_id: str) -> Optional[str]:
        """Export processed document to Word format"""
        try:
            session = self.get_session(session_id)
            if not session:
                return None
            
            # Get the populated content or original content
            content = session.get('populated_content', session['processed_content'])
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Church Outline with Bible Verses', 0)
            
            # Add content
            paragraphs = content.split('\\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            
            return temp_file.name
        
        except Exception as e:
            print(f"Error exporting to DOCX: {str(e)}")
            return None
    
    def get_verse_statistics(self) -> Dict:
        """Get statistics about the Bible database"""
        try:
            stats = self.bible_db.get_statistics()
            return {
                'success': True,
                'statistics': stats
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error getting statistics: {str(e)}'
            }

