import os
import tempfile
from typing import Dict, List, Optional
import PyPDF2
import pdfplumber
from docx import Document
from docx.shared import Inches
import uuid
import sqlite3
import json
from datetime import datetime
from src.utils.verse_parser import VerseParser
from src.utils.sqlite_bible_database import SQLiteBibleDatabase

class DocumentProcessor:
    def __init__(self):
        self.parser = VerseParser()
        self.bible_db = SQLiteBibleDatabase()
        self.supported_formats = ['.pdf', '.doc', '.docx']
        self.sessions_db_path = '/tmp/sessions.db'
        self._init_sessions_db()
    
    def _init_sessions_db(self):
        """Initialize sessions database"""
        conn = sqlite3.connect(self.sessions_db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_sessions (
                session_id TEXT PRIMARY KEY,
                original_filename TEXT,
                processed_content TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()
        conn.close()
    
    def process_uploaded_file(self, file_path: str, original_filename: str, text_content: str = None) -> Dict:
        """Process an uploaded file and extract text content"""
        try:
            if text_content is None:
                # Determine file type
                file_ext = os.path.splitext(original_filename)[1].lower()
                
                if file_ext not in self.supported_formats:
                    raise ValueError(f"Unsupported file format: {file_ext}")
                
                # Extract text based on file type
                if file_ext == '.pdf':
                    text_content = self._extract_pdf_text(file_path)
                elif file_ext in ['.doc', '.docx']:
                    text_content = self._extract_docx_text(file_path)
                else:
                    raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Create session
            session_id = str(uuid.uuid4())
            
            # Save to sessions database
            conn = sqlite3.connect(self.sessions_db_path)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO user_sessions (session_id, original_filename, processed_content)
                VALUES (?, ?, ?)
            ''', (session_id, original_filename, text_content))
            conn.commit()
            conn.close()
            
            # Extract verse references
            references = self.parser.extract_references_from_text(text_content)
            
            return {
                'success': True,
                'session_id': session_id,
                'original_filename': original_filename,
                'text_content': text_content,
                'detected_references': references,
                'reference_count': len(references)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_content = ""
        
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
        except Exception:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
            except Exception as e:
                raise Exception(f"Failed to extract text from PDF: {str(e)}")
        
        return text_content.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return text_content.strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def populate_document_with_verses(self, session_id: str, format_type: str = 'inline') -> Dict:
        """Populate document with verse content"""
        try:
            # Get session
            session_data = self.get_session(session_id)
            if not session_data:
                return {'success': False, 'error': 'Session not found'}
            
            text_content = session_data['processed_content']
            
            # Extract references
            references = self.parser.extract_references_from_text(text_content)
            
            # Look up verses
            all_verses = self.bible_db.lookup_verses_by_references(references)
            
            # Create populated content
            populated_content = self._format_content_with_verses(
                text_content, references, all_verses, format_type
            )
            
            # Update session
            self.update_session_content(session_id, populated_content)
            
            return {
                'success': True,
                'session_id': session_id,
                'populated_content': populated_content,
                'references': references,
                'verses': all_verses,
                'format_type': format_type
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _format_content_with_verses(self, content: str, references: List[str], 
                                   verses: List[Dict], format_type: str) -> str:
        """Format content with verse text"""
        # Create verse mapping
        verse_map = {}
        for verse in verses:
            ref = verse.get('original_reference', verse['reference'])
            if ref not in verse_map:
                verse_map[ref] = []
            verse_map[ref].append(verse)
        
        populated_content = content
        
        if format_type == 'inline':
            # Replace references with reference + verse text
            for ref in references:
                if ref in verse_map:
                    verses_text = []
                    for verse in verse_map[ref]:
                        verses_text.append(f"({verse['text']})")
                    replacement = f"{ref} {' '.join(verses_text)}"
                    populated_content = populated_content.replace(ref, replacement)
        
        elif format_type == 'footnotes':
            # Add footnotes at the end
            footnotes = []
            footnote_num = 1
            
            for ref in references:
                if ref in verse_map:
                    # Replace reference with footnote marker
                    populated_content = populated_content.replace(
                        ref, f"{ref}[{footnote_num}]"
                    )
                    
                    # Add footnote
                    for verse in verse_map[ref]:
                        footnotes.append(f"[{footnote_num}] {verse['reference']}: {verse['text']}")
                        footnote_num += 1
            
            if footnotes:
                populated_content += "\n\n--- Footnotes ---\n" + "\n".join(footnotes)
        
        return populated_content
    
    def export_to_docx(self, session_id: str) -> Optional[str]:
        """Export session content to Word document"""
        try:
            # Get session
            session_data = self.get_session(session_id)
            if not session_data:
                return None
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Church Outline with Bible Verses', 0)
            
            # Add content
            content = session_data['processed_content']
            paragraphs = content.split('\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Detect if this is a heading (starts with Roman numerals or letters)
                    if self._is_heading(paragraph_text):
                        doc.add_heading(paragraph_text.strip(), level=1)
                    elif self._is_subheading(paragraph_text):
                        doc.add_heading(paragraph_text.strip(), level=2)
                    else:
                        doc.add_paragraph(paragraph_text.strip())
            
            # Save to temporary file
            temp_dir = tempfile.gettempdir()
            filename = f"outline_{session_id}.docx"
            file_path = os.path.join(temp_dir, filename)
            doc.save(file_path)
            
            return file_path
            
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return None
    
    def _is_heading(self, text: str) -> bool:
        """Check if text is a main heading"""
        text = text.strip()
        # Check for Roman numerals at start
        return bool(text and (
            text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.')) or
            text.startswith('Message ') or
            text.startswith('Scripture Reading:')
        ))
    
    def _is_subheading(self, text: str) -> bool:
        """Check if text is a subheading"""
        text = text.strip()
        # Check for letters at start
        return bool(text and text.startswith(('A.', 'B.', 'C.', 'D.', 'E.', 'F.', 'G.', 'H.')))
    
    def get_session(self, session_id: str) -> Optional[Dict]:
        """Get session data"""
        try:
            conn = sqlite3.connect(self.sessions_db_path)
            cursor = conn.cursor()
            cursor.execute('''
                SELECT session_id, original_filename, processed_content, created_at, updated_at
                FROM user_sessions WHERE session_id = ?
            ''', (session_id,))
            
            result = cursor.fetchone()
            conn.close()
            
            if result:
                return {
                    'session_id': result[0],
                    'original_filename': result[1],
                    'processed_content': result[2],
                    'created_at': result[3],
                    'updated_at': result[4]
                }
            return None
        except Exception as e:
            print(f"Error getting session: {e}")
            return None
    
    def update_session_content(self, session_id: str, content: str) -> bool:
        """Update session content"""
        try:
            conn = sqlite3.connect(self.sessions_db_path)
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE user_sessions 
                SET processed_content = ?, updated_at = CURRENT_TIMESTAMP
                WHERE session_id = ?
            ''', (content, session_id))
            
            success = cursor.rowcount > 0
            conn.commit()
            conn.close()
            return success
        except Exception as e:
            print(f"Error updating session: {e}")
            return False

